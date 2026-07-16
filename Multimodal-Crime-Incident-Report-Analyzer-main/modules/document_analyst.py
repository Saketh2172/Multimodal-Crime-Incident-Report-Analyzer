"""
Multimodal Crime/Incident Report Analyzer — document pipeline (police / official PDFs).

Extracts text from PDFs (pdfplumber + PyMuPDF; OCR fallback for scans), Word (.docx),
and plain .txt under data/documents/. Structured fields match the course brief:
Report_ID, Incident_Type, Date, Location, Officer, Summary (+ Suspect_Description, Outcome).
"""

from __future__ import annotations

import io
import re
import sys
import warnings
from pathlib import Path

import pandas as pd

_MODULE_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _MODULE_DIR.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

_DATA_DOCS = _PROJECT_ROOT / "data" / "documents"
_OUTPUTS = _PROJECT_ROOT / "outputs"
_DOC_EXTENSIONS = {".pdf", ".docx", ".txt"}

_KEYWORD_INCIDENT: tuple[tuple[tuple[str, ...], str], ...] = (
    (("homicide", "murder", "manslaughter", "fatal shooting"), "homicide / assault"),
    (("shooting", "shot", "gunfire", "weapon", "firearm"), "weapons / shooting"),
    (("stabb", "assault", "battery", "attack"), "assault"),
    (("robbery", "burglary", "theft", "stolen", "larceny"), "theft / robbery"),
    (("arson", "fire", "smoke", "burning"), "fire / arson"),
    (("traffic", "vehicle crash", "collision", "crash", "dwi", "dui"), "traffic / accident"),
    (("riot", "disturbance", "disorderly", "fight", "domestic"), "public disturbance"),
    (("1033", "training plan", "proposal", "equipment request"), "administrative / training"),
    (("foia", "memorandum", "acquisition"), "administrative / training"),
)

_OFFICER_TITLE_PATTERN = re.compile(
    r"(?i)(?:officer|det\.|detective|chief|sergeant|sgt\.|"
    r"lieutenant|lt\.|captain|cpt\.|deputy|marshal)\s+"
)

_NAME_TOKEN = re.compile(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2}(?=\s|[,.]|$)")

_OUTCOME_PHRASES = (
    (("arrest", "arrested", "taken into custody", "booked"), "arrest / custody"),
    (("closed", "cleared", "concluded", "resolved"), "closed / resolved"),
    (("pending", "ongoing", "under investigation", "active investigation"), "pending investigation"),
    (("dismiss", "dropped", "no charges", "declined prosecution"), "dismissed / no charges"),
    (("fatal", "deceased", "died"), "fatal outcome noted"),
    (("recovered", "evidence collected"), "recovery / evidence"),
)

_SUSPIC_HINT = re.compile(
    r"(?i)(?:suspect|subject|perpetrator|accused|offender)\s*[:\-]?\s*(.{3,120}?)(?:\.|$|;|\n)",
)

_DATE_LINE_PATTERN = re.compile(
    r"(?i)(?:(?:date\s+)?reported|date|occurred)\s*[:\-]\s*"
    r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}\b)"
)

_LOCATION_LINE_PATTERN = re.compile(
    r"(?i)^\s*location\s*[:\-]\s*(.+?)\s*$",
    re.MULTILINE,
)

_NLP_CHAR_LIMIT = 1_000_000

_OUTPUT_COLUMNS = [
    "Report_ID",
    "Incident_Type",
    "Date",
    "Location",
    "Officer",
    "Summary",
    "Suspect_Description",
    "Outcome",
]


def _ensure_directories() -> None:
    _DATA_DOCS.mkdir(parents=True, exist_ok=True)
    _OUTPUTS.mkdir(parents=True, exist_ok=True)


def _list_document_files() -> list[Path]:
    if not _DATA_DOCS.is_dir():
        return []
    out: list[Path] = []
    for p in _DATA_DOCS.iterdir():
        if p.is_file() and not p.name.startswith(".") and p.suffix.lower() in _DOC_EXTENSIONS:
            out.append(p)
    return sorted(out, key=lambda x: x.name.lower())


def _extract_pdf_native(path: Path) -> str:
    chunks: list[str] = []
    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    chunks.append(t.strip())
    except Exception as e:
        warnings.warn(f"pdfplumber: {path.name}: {e}", stacklevel=2)

    text = "\n".join(chunks).strip()
    if len(text) < 80:
        try:
            import fitz

            doc = fitz.open(str(path))
            alt: list[str] = []
            for i in range(len(doc)):
                t = doc.load_page(i).get_text()
                if t and t.strip():
                    alt.append(t.strip())
            doc.close()
            merged = "\n".join(alt).strip()
            if len(merged) > len(text):
                text = merged
        except Exception as e:
            warnings.warn(f"PyMuPDF text: {path.name}: {e}", stacklevel=2)
    return text


def _ocr_pdf_pages(path: Path, max_pages: int = 25) -> str:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except Exception as e:
        warnings.warn(f"OCR skipped (need PyMuPDF, pytesseract, Pillow, Tesseract binary): {e}", stacklevel=2)
        return ""
    parts: list[str] = []
    try:
        doc = fitz.open(str(path))
        n = min(len(doc), max_pages)
        for i in range(n):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            t = pytesseract.image_to_string(img)
            if t and t.strip():
                parts.append(t.strip())
        doc.close()
    except Exception as e:
        warnings.warn(f"OCR failed for {path.name}: {e}", stacklevel=2)
        return ""
    return "\n".join(parts).strip()


def _extract_text_pdf(path: Path) -> str:
    text = _extract_pdf_native(path)
    if len(text) < 80:
        ocr_text = _ocr_pdf_pages(path)
        if len(ocr_text) > len(text):
            text = ocr_text
    return text


def _extract_text_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras).strip()


def _extract_text(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".txt":
        return path.read_text(encoding="utf-8", errors="replace").strip()
    if suf == ".pdf":
        return _extract_text_pdf(path)
    if suf == ".docx":
        return _extract_text_docx(path)
    return ""


def _normalize_ws(s: str) -> str:
    return " ".join((s or "").split())


def _incident_type_from_text(lowered: str) -> str:
    for keys, label in _KEYWORD_INCIDENT:
        if any(k in lowered for k in keys):
            return label
    if "police" in lowered or "department" in lowered or "sheriff" in lowered:
        return "law enforcement report"
    return "other / unspecified"


def _officer_from_text(text: str) -> str:
    for m in _OFFICER_TITLE_PATTERN.finditer(text):
        tail = text[m.end() :].strip()
        nm = _NAME_TOKEN.match(tail)
        if nm:
            return _normalize_ws(nm.group(0))
    return "N/A"


def _outcome_from_text(lowered: str) -> str:
    for keys, label in _OUTCOME_PHRASES:
        if any(k in lowered for k in keys):
            return label
    return "N/A"


def _suspect_from_text(text: str) -> str:
    m = _SUSPIC_HINT.search(text)
    if not m:
        return "N/A"
    s = _normalize_ws(m.group(1))
    s = re.sub(r"(?i)^described as\s+", "", s).strip()
    return s if s else "N/A"


def _dates_from_spacy(doc) -> str:
    dates = [e.text.strip() for e in doc.ents if e.label_ == "DATE" and e.text.strip()]
    if not dates:
        return "N/A"
    seen: set[str] = set()
    uniq: list[str] = []
    for d in dates:
        key = d.lower()
        if key not in seen and len(d) < 80:
            seen.add(key)
            uniq.append(d)
        if len(uniq) >= 4:
            break
    return "; ".join(uniq) if uniq else "N/A"


def _locations_from_spacy(doc) -> str:
    want = {"GPE", "LOC", "FAC"}
    locs = [_normalize_ws(e.text) for e in doc.ents if e.label_ in want and e.text.strip()]
    if not locs:
        return "N/A"
    seen: set[str] = set()
    out: list[str] = []
    for L in locs:
        k = L.lower()
        if k not in seen and len(L) < 120:
            seen.add(k)
            out.append(L)
        if len(out) >= 5:
            break
    return "; ".join(out) if out else "N/A"


def _officer_from_spacy(doc, text: str) -> str:
    static = _officer_from_text(text)
    if static != "N/A":
        return static
    for ent in doc.ents:
        if ent.label_ != "PERSON":
            continue
        start = max(0, ent.start_char - 48)
        ctx = text[start : ent.start_char].lower()
        if any(w in ctx for w in ("officer", "detective", "chief", "sergeant", "deputy", "lt.", "capt.")):
            return _normalize_ws(ent.text)
    return "N/A"


def _summary_from_body(body: str, max_chars: int = 600) -> str:
    s = _normalize_ws(body)
    if not s:
        return "N/A"
    if len(s) <= max_chars:
        return s
    return s[: max_chars - 3].rsplit(" ", 1)[0] + "..."


_ANY_DEPT_PAT = re.compile(
    r"([\w\s,\.'-]{3,60}?(?:Police Department|Sheriff(?:'s\s*Office)?|"
    r"Department of Public Safety|Police Dept\.))",
    re.IGNORECASE,
)
_COVER_FROM_PAT = re.compile(r"(?im)^From:\s*(.+)$")
_COVER_DATE_PAT = re.compile(r"(?im)^Date[W\s]?:\s*(.+)$")


_CLEAN_DEPT_PAT = re.compile(
    r"([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*\s+(?:Police Department|Sheriff(?:'s\s*Office)?"
    r"|Department of Public Safety|Police Dept\.))",
)


def _extract_dept_name(text: str) -> str:
    """Extract the most prominent department name from a block of text."""
    # Prefer names at the very start of the text (first 6 lines)
    for line in text.strip().splitlines()[:6]:
        m = _ANY_DEPT_PAT.search(line)
        if m:
            raw = m.group(1).strip()
            # Strip sentence fragments — keep only the proper-noun portion
            clean = _CLEAN_DEPT_PAT.search(raw)
            return clean.group(1).strip() if clean else raw
    # Fallback: scan full text
    m = _ANY_DEPT_PAT.search(text)
    if m:
        raw = m.group(1).strip()
        clean = _CLEAN_DEPT_PAT.search(raw)
        return clean.group(1).strip() if clean else raw
    return ""


def _split_pdf_sections(path: Path) -> list[dict[str, str]]:
    """Split a multi-department PDF into per-section dicts with dept, date, officer, text."""
    try:
        import pdfplumber
    except ImportError:
        return []

    sections: list[dict[str, str]] = []
    current_pages: list[str] = []
    current_dept: str = ""
    current_date: str = "N/A"
    current_officer: str = "N/A"

    def _flush():
        if current_pages and current_dept:
            body = "\n".join(current_pages).strip()
            if body:
                sections.append({
                    "dept": current_dept,
                    "date": current_date,
                    "officer": current_officer,
                    "body": body,
                })

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if not text.strip():
                continue

            # Cover letter: "To: Whom it may Concern" (also catch garbled OCR variants)
            is_cover = bool(
                re.search(r"(?i)whom\s+it\s+may\s+concern", text)
                and re.search(r"(?i)^From:", text, re.MULTILINE)
            )

            if is_cover:
                _flush()
                current_pages = [text]
                fm = _COVER_FROM_PAT.search(text)
                dm = _COVER_DATE_PAT.search(text)
                raw_from = fm.group(1).strip() if fm else ""
                current_date = dm.group(1).strip() if dm else "N/A"

                # If From: is a person title (e.g. "Lt. Brett Hibbs"), find dept in body
                if raw_from and not _ANY_DEPT_PAT.search(raw_from):
                    current_officer = raw_from
                    dept = _extract_dept_name(text)
                    current_dept = dept if dept else raw_from
                else:
                    dept = _extract_dept_name(raw_from) or _extract_dept_name(text)
                    current_dept = dept if dept else raw_from
                    current_officer = "N/A"
            else:
                # Detect a new section by department name change (no cover letter)
                page_dept = _extract_dept_name(text)
                if page_dept and current_dept and page_dept.lower() != current_dept.lower():
                    # Only start new section if the first few lines strongly indicate a new dept
                    first_lines = "\n".join(text.strip().splitlines()[:3])
                    if _ANY_DEPT_PAT.search(first_lines):
                        _flush()
                        current_pages = [text]
                        current_dept = page_dept
                        current_date = "N/A"
                        current_officer = "N/A"
                        continue
                elif not current_dept and page_dept:
                    # Very first section in file (no cover letter at all)
                    current_dept = page_dept

                current_pages.append(text)

    _flush()
    return sections


def _analyze_document(body: str, filename: str, override_date: str = "",
                      override_location: str = "", override_officer: str = "") -> dict[str, str]:
    lowered = body.lower()
    nlp = None
    doc = None
    try:
        import spacy

        nlp = spacy.load("en_core_web_sm")
        doc = nlp(body[:_NLP_CHAR_LIMIT])
    except Exception as e:
        warnings.warn(f"spaCy NER unavailable ({e}); using regex/heuristics only.", stacklevel=2)

    incident = _incident_type_from_text(lowered)
    summary = _summary_from_body(body) if body else f"(no extractable text from {filename})"

    if override_date:
        dt = override_date
    elif doc is not None:
        dt = _dates_from_spacy(doc)
    else:
        dt = "N/A"
        m = _DATE_LINE_PATTERN.search(body)
        if m:
            dt = m.group(1).strip()

    if override_location:
        loc = override_location
    elif doc is not None:
        loc = _locations_from_spacy(doc)
    else:
        loc = "N/A"
        ml = _LOCATION_LINE_PATTERN.search(body)
        if ml:
            loc = _normalize_ws(ml.group(1))[:200] or "N/A"

    if override_officer:
        off = override_officer
    elif doc is not None:
        off = _officer_from_spacy(doc, body)
    else:
        off = _officer_from_text(body)

    if dt == "N/A" and body and not override_date:
        m = _DATE_LINE_PATTERN.search(body)
        if m:
            dt = m.group(1).strip()

    if loc == "N/A" and body and not override_location:
        ml = _LOCATION_LINE_PATTERN.search(body)
        if ml:
            loc = _normalize_ws(ml.group(1))[:200] or "N/A"

    suspect = _suspect_from_text(body)
    outcome = _outcome_from_text(lowered)

    return {
        "Incident_Type": incident,
        "Date": dt,
        "Location": loc,
        "Officer": off,
        "Summary": summary,
        "Suspect_Description": suspect,
        "Outcome": outcome,
    }


def run_document_pipeline() -> None:
    _ensure_directories()
    paths = _list_document_files()

    rows: list[dict[str, object]] = []
    doc_counter = 1

    for path in paths:
        if path.suffix.lower() == ".pdf":
            sections = _split_pdf_sections(path)
            if sections:
                for sec in sections:
                    report_id = f"DOC-{doc_counter:03d}"
                    doc_counter += 1
                    fields = _analyze_document(
                        sec["body"],
                        path.name,
                        override_date=sec["date"],
                        override_location=sec["dept"],
                        override_officer=sec["officer"],
                    )
                    rows.append({"Report_ID": report_id, **fields})
                continue

        # Non-PDF or PDF with no detectable sections → one row per file
        report_id = f"DOC-{doc_counter:03d}"
        doc_counter += 1
        try:
            body = _extract_text(path)
        except Exception as e:
            warnings.warn(f"Could not read {path.name}: {e}", stacklevel=2)
            body = ""
        fields = _analyze_document(body, path.name)
        rows.append({"Report_ID": report_id, **fields})

    out_df = pd.DataFrame(rows)
    if out_df.empty:
        out_df = pd.DataFrame(columns=_OUTPUT_COLUMNS)
    else:
        out_df = out_df[_OUTPUT_COLUMNS]

    out_csv = _OUTPUTS / "document_output.csv"
    out_df.to_csv(out_csv, index=False)
    if paths:
        print(out_df.to_string(index=False))
        print(f"[Document Analyst] Processed {len(rows)} section(s) from {len(paths)} file(s) → {out_csv.name}", flush=True)
    else:
        print(
            f"[Document Analyst] No files in {_DATA_DOCS.relative_to(_PROJECT_ROOT)} "
            f"(supported: {', '.join(sorted(_DOC_EXTENSIONS))})",
            flush=True,
        )
        print(f"[Document Analyst] Wrote empty {out_csv.name}", flush=True)


if __name__ == "__main__":
    run_document_pipeline()
